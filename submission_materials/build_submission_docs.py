from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "outputs"
OUTPUT.mkdir(parents=True, exist_ok=True)


def register_chinese_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    for path in candidates:
        if path.exists():
            name = "SubmissionCN"
            pdfmetrics.registerFont(TTFont(name, str(path)))
            return name
    raise FileNotFoundError("未找到可用的中文字体文件。")


FONT_NAME = register_chinese_font()


def set_doc_fonts(document: Document) -> None:
    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "List Bullet"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)


def build_docx(title: str, subtitle: str, sections: list[dict], path: Path) -> None:
    doc = Document()
    set_doc_fonts(doc)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Microsoft YaHei"
    title_p._element.rPr = title_p._element.get_or_add_pPr()

    if subtitle:
      sub = doc.add_paragraph()
      sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
      sub_run = sub.add_run(subtitle)
      sub_run.font.size = Pt(11)
      sub_run.font.name = "Microsoft YaHei"

    doc.add_paragraph("")
    for section in sections:
        doc.add_heading(section["title"], level=1)
        for paragraph in section.get("paragraphs", []):
            p = doc.add_paragraph(paragraph)
            p.style = doc.styles["Normal"]
        for bullet in section.get("bullets", []):
            p = doc.add_paragraph(bullet, style="List Bullet")
        for sub in section.get("subsections", []):
            doc.add_heading(sub["title"], level=2)
            for paragraph in sub.get("paragraphs", []):
                doc.add_paragraph(paragraph)
            for bullet in sub.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")
    doc.save(path)


def build_pdf(title: str, subtitle: str, sections: list[dict], path: Path) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName=FONT_NAME,
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#10233d"),
    )
    sub_style = ParagraphStyle(
        "SubCN",
        parent=styles["Normal"],
        fontName=FONT_NAME,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#5b6b7f"),
    )
    h1 = ParagraphStyle(
        "H1CN",
        parent=styles["Heading1"],
        fontName=FONT_NAME,
        fontSize=15,
        leading=22,
        textColor=colors.HexColor("#10233d"),
        spaceBefore=10,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2CN",
        parent=styles["Heading2"],
        fontName=FONT_NAME,
        fontSize=12,
        leading=18,
        textColor=colors.HexColor("#1f3b63"),
        spaceBefore=6,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=FONT_NAME,
        fontSize=10.5,
        leading=17,
        textColor=colors.black,
        spaceAfter=6,
    )
    bullet = ParagraphStyle(
        "BulletCN",
        parent=body,
        leftIndent=14,
        firstLineIndent=-10,
    )

    story: list = [Paragraph(title, title_style)]
    if subtitle:
        story.extend([Spacer(1, 0.15 * cm), Paragraph(subtitle, sub_style)])
    story.append(Spacer(1, 0.4 * cm))

    for section in sections:
        story.append(Paragraph(section["title"], h1))
        for paragraph in section.get("paragraphs", []):
            story.append(Paragraph(paragraph, body))
        for item in section.get("bullets", []):
            story.append(Paragraph(f"• {item}", bullet))
        for sub in section.get("subsections", []):
            story.append(Paragraph(sub["title"], h2))
            for paragraph in sub.get("paragraphs", []):
                story.append(Paragraph(paragraph, body))
            for item in sub.get("bullets", []):
                story.append(Paragraph(f"• {item}", bullet))
        story.append(Spacer(1, 0.15 * cm))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    doc.build(story)


def write_markdown(path: Path, title: str, sections: Iterable[dict]) -> None:
    lines = [f"# {title}", ""]
    for section in sections:
        lines.extend([f"## {section['title']}", ""])
        for paragraph in section.get("paragraphs", []):
            lines.extend([paragraph, ""])
        for bullet in section.get("bullets", []):
            lines.append(f"- {bullet}")
        if section.get("bullets"):
            lines.append("")
        for sub in section.get("subsections", []):
            lines.extend([f"### {sub['title']}", ""])
            for paragraph in sub.get("paragraphs", []):
                lines.extend([paragraph, ""])
            for bullet in sub.get("bullets", []):
                lines.append(f"- {bullet}")
            if sub.get("bullets"):
                lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


overview_sections = [
    {
        "title": "前言",
        "paragraphs": [
            "知行雷达是一套面向高校学生成长支持场景的行为分析与预警系统，围绕“多源数据融合、学生画像构建、风险识别、个性化建议与移动化服务”形成完整闭环。",
            "项目聚焦高校在学生管理中面临的数据分散、识别滞后、干预缺少依据和成果展示不足等问题，提供可落地、可展示、可扩展的数字化解决方案。"
        ]
    },
    {
        "title": "创意描述",
        "paragraphs": [
            "系统将上网、门禁、图书馆、课程行为、体测、英语、奖学金、竞赛等多源数据统一建模，不仅输出风险结果，还同步生成学生个体画像、群体画像、预测结果、图表分析和智能助手交互。"
        ],
        "bullets": [
            "构建学生个体画像与群体画像双体系",
            "识别不少于4类学生模式并细分子类",
            "形成8张以上全样本分析成果图",
            "在Web端与安卓端提供一致的数据服务能力"
        ]
    },
    {
        "title": "功能简介",
        "bullets": [
            "管理端：总览、风险名单、院系对比、干预工作台、预测模块、分析成果、学生详情与完整报告",
            "学生端：首页、我的画像、在线预测、群体对比、个性化报告、全样本分析、智能助手、设置中心",
            "移动端：学生端与管理端功能同步，并提供统一品牌与轻量化交互"
        ]
    },
    {
        "title": "特色综述",
        "bullets": [
            "多模型联合：风险预测、奖学金概率、综合成绩预测、健康水平分类、变化趋势预测、学习投入分类、综合发展分类、英语四级与六级通过概率",
            "解释能力强：从学习投入、行为规律、健康发展、资源利用等多个维度解释结果",
            "产品化程度高：同时具备后台管理、学生自助查看、移动端访问与智能问答能力",
            "工程链路完整：前后端分离、数据库、部署脚本、接口契约、模型文件、图表资源齐备"
        ]
    },
    {
        "title": "开发工具与技术",
        "bullets": [
            "后端：Flask + SQLite + Pandas + Scikit-learn",
            "前端：Vue 3 + TypeScript + uni-app",
            "可视化：ECharts + 系统导出图表",
            "智能助手：千问大模型接口接入"
        ]
    },
    {
        "title": "应用对象",
        "bullets": [
            "学校管理人员",
            "学院管理人员与辅导员",
            "学生个体"
        ]
    },
    {
        "title": "应用环境",
        "paragraphs": [
            "系统支持本地开发环境、局域网演示环境与服务器部署环境，可通过浏览器访问Web端，也可通过安卓安装包使用移动端。"
        ]
    },
    {
        "title": "结语",
        "paragraphs": [
            "知行雷达并非只停留在模型实验层面，而是完成了从数据接入、模型训练、结果解释、产品展示到移动化服务的整体落地，为高校学生管理提供更及时、更可视、更可执行的数字化支持。"
        ]
    }
]

solution_sections = [
    {
        "title": "项目背景与需求理解",
        "paragraphs": [
            "高校学生管理工作正在从经验驱动向数据驱动转变，但现实中仍存在数据源分散、预警滞后、人工筛查成本高、干预依据不清晰、缺少统一展示入口等问题。",
            "本项目结合赛题要求，面向“学生行为分析与干预”场景，目标是搭建一套可运行的分析平台，并提供可视化成果、画像结果、风险预警与个性化服务。"
        ],
        "bullets": [
            "完成多源学生行为数据整合",
            "识别学生风险与成长差异",
            "形成个体画像与群体画像",
            "产出8~10个行为分析成果",
            "支持Web端与移动端展示"
        ]
    },
    {
        "title": "系统总体方案",
        "paragraphs": [
            "系统采用前后端分离架构，后端承担统一鉴权、数据读取、模型推理、报告生成与图表服务；Web端与移动端均通过统一API访问。"
        ],
        "bullets": [
            "后端：Flask 提供统一业务接口",
            "管理端Web：风险预警、群体分析、学生详情、模型结果查看",
            "学生端Web：画像、预测、报告、建议、智能助手",
            "移动端：学生端与管理端核心功能同步"
        ]
    },
    {
        "title": "数据来源与处理流程",
        "bullets": [
            "学生基本信息",
            "上网行为数据",
            "门禁与图书馆数据",
            "课程行为、作业、测验与考试数据",
            "英语、奖学金、竞赛、体测等专项数据"
        ],
        "subsections": [
            {
                "title": "数据治理流程",
                "bullets": [
                    "统一学号主键并对齐基础信息",
                    "对原始缺失与真实0值进行区分处理",
                    "生成基础特征与进阶特征两层结构",
                    "形成训练主表与展示主表"
                ]
            }
        ]
    },
    {
        "title": "学生画像与群体画像设计",
        "paragraphs": [
            "系统通过聚类与规则解释相结合，构建4类主画像，并在每一类下继续细分子类。"
        ],
        "bullets": [
            "高投入稳健型",
            "低投入风险型",
            "夜间波动型",
            "发展过渡型"
        ]
    },
    {
        "title": "预测模型体系",
        "paragraphs": [
            "系统围绕学生风险与成长表现建立多任务预测体系，不再局限于单一风险分数。"
        ],
        "bullets": [
            "风险预测模型 AUC 约 0.92",
            "英语四级通过预测 AUC 约 0.86",
            "英语六级通过预测 AUC 约 0.99",
            "健康分类 Macro F1 约 0.75",
            "学习投入分类 Macro F1 约 0.92",
            "综合发展分类 Macro F1 约 0.90"
        ]
    },
    {
        "title": "可解释性与干预机制",
        "paragraphs": [
            "系统在结果展示中同步输出关键影响因素、解释步骤、干预建议和报告摘要，避免只有分数没有依据。"
        ],
        "bullets": [
            "关键特征贡献展示",
            "多维度画像解释",
            "预测结果原因说明",
            "干预建议与行动方向"
        ]
    },
    {
        "title": "功能模块设计",
        "subsections": [
            {
                "title": "管理端",
                "bullets": [
                    "总览大屏",
                    "风险名单",
                    "院系对比",
                    "干预工作台",
                    "预测模块",
                    "分析成果",
                    "学生详情与完整报告"
                ]
            },
            {
                "title": "学生端",
                "bullets": [
                    "首页总览",
                    "我的画像",
                    "在线预测",
                    "群体对比",
                    "个性化报告",
                    "全样本分析",
                    "智能助手与设置中心"
                ]
            }
        ]
    },
    {
        "title": "移动端设计",
        "paragraphs": [
            "项目基于 uni-app 构建安卓端，在保留Web端核心能力的基础上，按移动端使用习惯重排界面，统一采用“知行雷达”品牌与轻蓝科技感风格。"
        ]
    },
    {
        "title": "系统实现与工程化",
        "bullets": [
            "统一数据库路径与账号体系，保证Web端与移动端共用同一套账号数据",
            "提供环境变量、部署说明、接口契约与运行脚本",
            "支持局域网调试、Docker部署与服务器部署扩展",
            "支持智能助手与外部大模型接口接入"
        ]
    },
    {
        "title": "创新点与应用价值",
        "bullets": [
            "从模型实验扩展到可落地的产品系统",
            "多任务预测与画像解释联动",
            "同一套数据服务Web端、学生端与移动端",
            "引入智能助手增强交互与指导能力"
        ]
    },
    {
        "title": "可行性分析与后续规划",
        "paragraphs": [
            "项目已经完成主要功能闭环，后续可继续扩展到更大规模数据接入、更多高校适配、干预流程闭环管理和服务器正式部署。"
        ]
    }
]

difficulty_sections = [
    {
        "title": "核心困难",
        "bullets": [
            "多源数据字段口径不一致，部分字段缺失值较多",
            "BMI、图书馆、门禁、竞赛等字段容易把缺失值误处理成真实0",
            "风险概率、风险标签与聚类画像之间存在口径不统一问题",
            "移动端早期存在页面乱码、路由配置损坏、基座白屏与旧缓存干扰",
            "Web端注册账号与移动端账号数据一度因为数据库路径错误而不一致",
            "部分模型指标过高或过低，需要识别是否存在泄漏或指标口径不合理"
        ]
    },
    {
        "title": "解决过程",
        "bullets": [
            "对基本特征逐项回查原始数据表，建立数据审计脚本，识别真实缺失与错误填充",
            "调整数据清洗逻辑，对关键字段不再一刀切填0，并在前后端明确显示“暂无原始记录”",
            "重新梳理风险口径，在页面中区分风险等级、画像类别与细分子类的语义",
            "重建移动端 pages.json、manifest.json、登录页与核心页面，统一为可运行 UTF-8 配置",
            "统一后端数据库路径到根目录数据库，重置测试账号密码，保证Web与移动端数据一致",
            "对多分类任务采用更合适的 Macro F1 / Accuracy 口径，并对高AUC任务说明类别不平衡背景"
        ]
    },
    {
        "title": "最终效果",
        "paragraphs": [
            "经过上述修正，系统已形成可运行、可展示、可解释的完整项目形态，同时具备管理端、学生端、安卓端和智能助手模块，为答辩与演示提供稳定支撑。"
        ]
    }
]

video_sections = [
    {
        "title": "视频总时长建议",
        "paragraphs": [
            "建议控制在 4 分 30 秒到 5 分钟之间，按照“背景—方案—系统演示—亮点总结”结构录制。"
        ]
    },
    {
        "title": "脚本分镜",
        "subsections": [
            {
                "title": "00:00 - 00:25 开场",
                "bullets": [
                    "介绍项目名称：知行雷达",
                    "说明项目面向高校学生行为分析与成长支持场景",
                    "点出多源数据、画像、预警、报告和移动端是核心亮点"
                ]
            },
            {
                "title": "00:25 - 00:55 方案概览",
                "bullets": [
                    "展示系统总体架构图",
                    "说明管理端、学生端、移动端和智能助手构成",
                    "简述多源数据和多任务预测体系"
                ]
            },
            {
                "title": "00:55 - 02:10 管理端演示",
                "bullets": [
                    "进入总览页，展示样本量、风险分布、注册覆盖与重点学生",
                    "进入风险名单，展示详情与完整报告入口",
                    "进入院系对比与预测模块，展示群体画像与单学生预测结果",
                    "进入分析成果，展示图表并打开图表解释"
                ]
            },
            {
                "title": "02:10 - 03:20 学生端演示",
                "bullets": [
                    "展示首页与我的画像",
                    "进入在线预测，演示多模型结果",
                    "展示群体对比与个性化报告",
                    "展示全样本分析与智能助手"
                ]
            },
            {
                "title": "03:20 - 04:10 安卓端演示",
                "bullets": [
                    "展示知行雷达安卓端首页",
                    "展示学生端关键模块",
                    "展示管理端关键模块"
                ]
            },
            {
                "title": "04:10 - 04:50 亮点总结",
                "bullets": [
                    "总结多源数据、画像解释、多任务预测、跨端一致和智能助手",
                    "强调系统具有实际落地与推广价值"
                ]
            }
        ]
    }
]

demo_sections = [
    {
        "title": "Demo 运行内容",
        "bullets": [
            "后端 Flask 服务",
            "管理端 Web 页面",
            "学生端 Web 页面",
            "安卓移动端安装包",
            "模型文件、图表文件与数据库文件"
        ]
    },
    {
        "title": "运行方式",
        "bullets": [
            "执行 Bootstrap-Dev.ps1 初始化环境",
            "执行 Start-Dev.ps1 启动前后端",
            "管理端访问 localhost:3200，学生端访问 localhost:3201，后端访问 localhost:5000",
            "安卓端可通过最新打包的 apk 安装体验"
        ]
    },
    {
        "title": "测试账号",
        "bullets": [
            "管理员：admin001 / 123456",
            "学生：200307700011 / 123456"
        ]
    },
    {
        "title": "提交建议",
        "paragraphs": [
            "最终提交时建议保留源码、数据库、模型文件、图表资源、环境说明和测试账号清单，确保评审可快速运行。"
        ]
    }
]

checklist_sections = [
    {
        "title": "必交材料",
        "bullets": [
            "项目概要介绍 PDF",
            "项目简介 PPT（不超过20页）",
            "项目详细方案 PDF",
            "项目演示视频 MP4",
            "企业要求其他材料"
        ]
    },
    {
        "title": "本次已生成",
        "bullets": [
            "项目概要介绍：docx + pdf + markdown",
            "项目详细方案：docx + pdf + markdown",
            "困难与解决过程：docx + pdf + markdown",
            "项目演示视频脚本：docx + pdf + markdown",
            "可运行Demo说明：docx + pdf + markdown",
            "20页项目简介PPT：pptx"
        ]
    },
    {
        "title": "仍需人工补充",
        "bullets": [
            "最终演示视频录制与压缩",
            "团队编号、团队名称、赛题编号等正式命名信息",
            "如有专利或软著，补充知识产权证明文件",
            "最终打包 ZIP 和网盘分享链接"
        ]
    }
]


DOCUMENTS = [
    ("01_项目概要介绍", "知行雷达项目概要介绍", "面向第十七届中国大学生服务外包创新创业大赛 A14 赛题", overview_sections),
    ("03_项目详细方案", "知行雷达项目详细方案", "面向第十七届中国大学生服务外包创新创业大赛 A14 赛题", solution_sections),
    ("05_困难与解决过程", "知行雷达困难与解决过程说明", "企业要求材料", difficulty_sections),
    ("04_项目演示视频脚本", "知行雷达项目演示视频脚本", "建议时长 4分30秒 - 5分钟", video_sections),
    ("06_可运行Demo说明", "知行雷达可运行 Demo 说明", "企业要求材料", demo_sections),
]


def main() -> None:
    for file_stem, title, subtitle, sections in DOCUMENTS:
        build_docx(title, subtitle, sections, OUTPUT / f"{file_stem}.docx")
        build_pdf(title, subtitle, sections, OUTPUT / f"{file_stem}.pdf")
        write_markdown(OUTPUT / f"{file_stem}.md", title, sections)
    write_markdown(OUTPUT / "00_提交清单.md", "知行雷达提交清单", checklist_sections)


if __name__ == "__main__":
    main()
