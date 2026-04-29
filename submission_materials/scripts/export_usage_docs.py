from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Preformatted


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "submission_materials" / "outputs" / "07_系统使用说明.md"
DOCX_OUT = ROOT / "submission_materials" / "outputs" / "07_系统使用说明_企业版.docx"
PDF_OUT = ROOT / "submission_materials" / "outputs" / "07_系统使用说明_企业版.pdf"


def read_markdown() -> list[str]:
    text = SOURCE.read_text(encoding="utf-8")
    return text.splitlines()


def configure_docx(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)


def export_docx(lines: list[str]) -> None:
    doc = Document()
    configure_docx(doc)

    in_code = False
    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            run = p.add_run(line)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(18)
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(14)
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(12)
            continue

        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\.\s*", "", line))
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
            continue

        if line.startswith("  - "):
            p = doc.add_paragraph(style="List Bullet 2")
            p.add_run(line[4:].strip())
            continue

        doc.add_paragraph(line)

    doc.save(DOCX_OUT)


def find_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\msyh.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    raise FileNotFoundError("No suitable Chinese font found.")


def export_pdf(lines: list[str]) -> None:
    font_path = find_font()
    pdfmetrics.registerFont(TTFont("ChineseFont", font_path))

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName="ChineseFont",
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h2_style = ParagraphStyle(
        "Heading2CN",
        parent=styles["Heading2"],
        fontName="ChineseFont",
        fontSize=14,
        leading=20,
        spaceBefore=8,
        spaceAfter=6,
    )
    h3_style = ParagraphStyle(
        "Heading3CN",
        parent=styles["Heading3"],
        fontName="ChineseFont",
        fontSize=12,
        leading=18,
        spaceBefore=6,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName="ChineseFont",
        fontSize=11,
        leading=18,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "CodeCN",
        parent=styles["Code"],
        fontName="ChineseFont",
        fontSize=9.5,
        leading=13,
    )

    story = []
    in_code = False
    bullet_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        if not bullet_buffer:
            return
        items = [ListItem(Paragraph(item, body_style)) for item in bullet_buffer]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
        story.append(Spacer(1, 4))
        bullet_buffer = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_bullets()
            in_code = not in_code
            continue

        if in_code:
            story.append(Preformatted(line, code_style))
            continue

        if not line.strip():
            flush_bullets()
            story.append(Spacer(1, 6))
            continue

        if line.startswith("- "):
            bullet_buffer.append(line[2:].strip())
            continue

        if line.startswith("  - "):
            bullet_buffer.append("    " + line[4:].strip())
            continue

        flush_bullets()

        if line.startswith("# "):
            story.append(Paragraph(line[2:].strip(), title_style))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:].strip(), h2_style))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), h3_style))
        else:
            story.append(Paragraph(line, body_style))

    flush_bullets()

    pdf = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    pdf.build(story)


if __name__ == "__main__":
    lines = read_markdown()
    export_docx(lines)
    export_pdf(lines)
    print(f"generated: {DOCX_OUT}")
    print(f"generated: {PDF_OUT}")
